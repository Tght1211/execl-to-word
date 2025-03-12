import sys
import os
import yaml
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                           QHBoxLayout, QPushButton, QLabel, QFileDialog, 
                           QTableWidget, QTableWidgetItem, QMessageBox,
                           QRadioButton, QButtonGroup, QProgressDialog)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont, QPalette, QColor
import pandas as pd
from docx import Document
import openpyxl

class ExcelWordConverter(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Excel转Word工具')
        self.setMinimumSize(800, 600)
        
        # 设置主窗口样式
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f0f2f5;
            }
            QPushButton {
                background-color: #1890ff;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #40a9ff;
            }
            QLabel {
                font-size: 14px;
                color: #262626;
            }
            QTableWidget {
                background-color: white;
                border: 1px solid #d9d9d9;
                border-radius: 4px;
            }
            QRadioButton {
                font-size: 14px;
                color: #262626;
                padding: 4px;
            }
        """)

        # 初始化变量
        self.excel_path = ''
        self.word_template_path = ''
        self.mapping_config = {}
        self.excel_data = None
        self.excel_wb = None
        
        # 创建主窗口部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        
        # 文件选择区域
        file_section = QWidget()
        file_layout = QVBoxLayout(file_section)
        
        # Excel文件选择
        excel_widget = QWidget()
        excel_layout = QHBoxLayout(excel_widget)
        self.excel_label = QLabel('Excel文件：未选择')
        excel_btn = QPushButton('选择Excel')
        excel_btn.clicked.connect(self.select_excel)
        excel_layout.addWidget(self.excel_label)
        excel_layout.addWidget(excel_btn)
        
        # Word模板选择
        word_widget = QWidget()
        word_layout = QHBoxLayout(word_widget)
        self.word_label = QLabel('Word模板：未选择')
        word_btn = QPushButton('选择模板')
        word_btn.clicked.connect(self.select_word_template)
        word_layout.addWidget(self.word_label)
        word_layout.addWidget(word_btn)
        
        file_layout.addWidget(excel_widget)
        file_layout.addWidget(word_widget)
        
        # 添加文档生成模式选择
        mode_widget = QWidget()
        mode_layout = QHBoxLayout(mode_widget)
        mode_label = QLabel('生成模式：')
        self.mode_group = QButtonGroup()
        self.separate_mode = QRadioButton('生成多个文档')
        self.separate_mode.setChecked(True)
        self.mode_group.addButton(self.separate_mode)
        mode_layout.addWidget(mode_label)
        mode_layout.addWidget(self.separate_mode)
        mode_layout.addStretch()
        
        # 映射配置表格
        self.mapping_table = QTableWidget(0, 4)
        self.mapping_table.setHorizontalHeaderLabels(['Excel单元格', '列标题预览', 'Word变量名', '操作'])
        self.mapping_table.setColumnWidth(0, 120)  # Excel单元格列宽
        self.mapping_table.setColumnWidth(1, 150)  # 列标题预览列宽
        self.mapping_table.setColumnWidth(2, 150)  # Word变量名列宽
        
        # 按钮区域
        button_widget = QWidget()
        button_layout = QHBoxLayout(button_widget)
        add_mapping_btn = QPushButton('添加映射')
        save_config_btn = QPushButton('保存配置')
        load_config_btn = QPushButton('加载配置')
        generate_btn = QPushButton('生成文档')
        
        add_mapping_btn.clicked.connect(self.add_mapping_row)
        save_config_btn.clicked.connect(self.save_config)
        load_config_btn.clicked.connect(self.load_config)
        generate_btn.clicked.connect(self.generate_documents)
        
        button_layout.addWidget(add_mapping_btn)
        button_layout.addWidget(save_config_btn)
        button_layout.addWidget(load_config_btn)
        button_layout.addWidget(generate_btn)
        
        # 添加所有部件到主布局
        layout.addWidget(file_section)
        layout.addWidget(mode_widget)
        layout.addWidget(self.mapping_table)
        layout.addWidget(button_widget)

    def select_excel(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择Excel文件", "", "Excel Files (*.xlsx *.xls)")
        if file_path:
            self.excel_path = file_path
            self.excel_label.setText(f'Excel文件：{os.path.basename(file_path)}')
            # 加载Excel文件
            try:
                self.excel_wb = openpyxl.load_workbook(file_path, data_only=True)
                self.excel_data = pd.read_excel(file_path)
            except Exception as e:
                QMessageBox.critical(self, '错误', f'加载Excel文件失败：{str(e)}')

    def get_cell_value(self, cell_ref):
        """获取Excel单元格的值"""
        try:
            if not self.excel_wb:
                return None
            sheet = self.excel_wb.active
            return sheet[cell_ref].value
        except Exception:
            return None

    def get_column_letter(self, cell_ref):
        """从单元格引用中获取列字母"""
        import re
        match = re.match(r'([A-Za-z]+)\d+', cell_ref)
        return match.group(1) if match else None

    def add_mapping_row(self):
        row_position = self.mapping_table.rowCount()
        self.mapping_table.insertRow(row_position)
        
        # 创建单元格编辑器
        cell_ref_item = QTableWidgetItem()
        self.mapping_table.setItem(row_position, 0, cell_ref_item)
        
        # 创建预览单元格
        preview_item = QTableWidgetItem()
        preview_item.setFlags(preview_item.flags() & ~Qt.ItemFlag.ItemIsEditable)  # 设置为只读
        self.mapping_table.setItem(row_position, 1, preview_item)
        
        # 创建变量名单元格
        var_item = QTableWidgetItem()
        self.mapping_table.setItem(row_position, 2, var_item)
        
        # 添加删除按钮
        delete_btn = QPushButton('删除')
        delete_btn.clicked.connect(lambda: self.mapping_table.removeRow(
            self.mapping_table.indexAt(delete_btn.pos()).row()))
        self.mapping_table.setCellWidget(row_position, 3, delete_btn)
        
        # 添加单元格编辑完成事件处理
        cell_ref_item.setData(Qt.ItemDataRole.UserRole, "cell_ref")
        
        # 连接单元格变化信号
        self.mapping_table.itemChanged.connect(self.on_cell_ref_changed)

    def on_cell_ref_changed(self, item):
        if item.data(Qt.ItemDataRole.UserRole) == "cell_ref":
            row = item.row()
            cell_ref = item.text().upper()  # 转换为大写
            preview_value = self.get_cell_value(cell_ref)
            if preview_value is not None:
                self.mapping_table.item(row, 1).setText(str(preview_value))
            else:
                self.mapping_table.item(row, 1).setText('无效的单元格引用')

    def select_word_template(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择Word模板", "", "Word Files (*.docx)")
        if file_path:
            self.word_template_path = file_path
            self.word_label.setText(f'Word模板：{os.path.basename(file_path)}')
            
    def save_config(self):
        if not self.mapping_table.rowCount():
            QMessageBox.warning(self, '警告', '没有可保存的映射配置！')
            return
            
        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存配置", "", "YAML Files (*.yaml)")
        if not file_path:
            return
            
        config = {
            'excel_path': self.excel_path,
            'word_template_path': self.word_template_path,
            'mappings': []
        }
        
        for row in range(self.mapping_table.rowCount()):
            cell_ref = self.mapping_table.item(row, 0)
            word_var = self.mapping_table.item(row, 2)
            if cell_ref and word_var:
                config['mappings'].append({
                    'cell_ref': cell_ref.text(),
                    'word_variable': word_var.text()
                })
                
        with open(file_path, 'w', encoding='utf-8') as f:
            yaml.dump(config, f, allow_unicode=True)
            
        QMessageBox.information(self, '成功', '配置已保存！')

    def load_config(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "加载配置", "", "YAML Files (*.yaml)")
        if not file_path:
            return
            
        with open(file_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
            
        self.excel_path = config.get('excel_path', '')
        self.word_template_path = config.get('word_template_path', '')
        
        if self.excel_path:
            self.excel_label.setText(f'Excel文件：{os.path.basename(self.excel_path)}')
            try:
                self.excel_wb = openpyxl.load_workbook(self.excel_path, data_only=True)
                self.excel_data = pd.read_excel(self.excel_path)
            except Exception:
                pass
                
        if self.word_template_path:
            self.word_label.setText(f'Word模板：{os.path.basename(self.word_template_path)}')
        
        # 清空现有映射
        self.mapping_table.setRowCount(0)
        
        # 添加配置中的映射
        for mapping in config.get('mappings', []):
            row_position = self.mapping_table.rowCount()
            self.mapping_table.insertRow(row_position)
            
            cell_ref_item = QTableWidgetItem(mapping['cell_ref'])
            cell_ref_item.setData(Qt.ItemDataRole.UserRole, "cell_ref")
            self.mapping_table.setItem(row_position, 0, cell_ref_item)
            
            preview_value = self.get_cell_value(mapping['cell_ref'])
            preview_item = QTableWidgetItem(str(preview_value) if preview_value is not None else '无效的单元格引用')
            preview_item.setFlags(preview_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.mapping_table.setItem(row_position, 1, preview_item)
            
            self.mapping_table.setItem(row_position, 2, QTableWidgetItem(mapping['word_variable']))
            
            delete_btn = QPushButton('删除')
            delete_btn.clicked.connect(lambda: self.mapping_table.removeRow(
                self.mapping_table.indexAt(delete_btn.pos()).row()))
            self.mapping_table.setCellWidget(row_position, 3, delete_btn)

    def generate_documents(self):
        if not self.excel_path or not self.word_template_path:
            QMessageBox.warning(self, '警告', '请先选择Excel文件和Word模板！')
            return
            
        if not self.mapping_table.rowCount():
            QMessageBox.warning(self, '警告', '请先配置映射关系！')
            return
            
        try:
            # 获取映射关系
            mappings = {}
            cell_refs = {}  # 存储单元格引用
            
            for row in range(self.mapping_table.rowCount()):
                cell_ref = self.mapping_table.item(row, 0)
                word_var = self.mapping_table.item(row, 2)
                if cell_ref and word_var:
                    var_name = word_var.text()
                    cell_ref_text = cell_ref.text().upper()
                    
                    # 标准化变量名格式
                    if not var_name.startswith('${'):
                        var_name = '${' + var_name
                    if not var_name.endswith('}'):
                        var_name = var_name + '}'
                    
                    mappings[var_name] = cell_ref_text
                    cell_refs[var_name] = cell_ref_text
            
            # 选择保存目录
            save_dir = QFileDialog.getExistingDirectory(self, "选择保存目录")
            if not save_dir:
                return

            # 使用openpyxl读取Excel数据
            wb = openpyxl.load_workbook(self.excel_path, data_only=True)
            sheet = wb.active
            
            # 获取数据行范围
            data_rows = []
            for row in range(4, sheet.max_row + 1):  # 从第4行开始（跳过标题行）
                if any(sheet.cell(row=row, column=1).value is not None for column in range(1, sheet.max_column + 1)):
                    data_rows.append(row)

            # 创建进度对话框
            progress = QProgressDialog("正在生成文档...", "取消", 0, len(data_rows), self)
            progress.setWindowTitle("处理中")
            progress.setWindowModality(Qt.WindowModality.WindowModal)
            progress.setMinimumDuration(0)
            progress.setValue(0)

            if self.separate_mode.isChecked():
                # 生成多个文档
                for idx, row_num in enumerate(data_rows, 1):
                    if progress.wasCanceled():
                        break
                    doc = Document(self.word_template_path)
                    self._process_document(doc, sheet, row_num, mappings)
                    output_path = os.path.join(save_dir, f'生成文档_{idx}.docx')
                    doc.save(output_path)
                    progress.setValue(idx)
                    QApplication.processEvents()
            
            progress.close()
            if not progress.wasCanceled():
                QMessageBox.information(self, '成功', '文档生成完成！')
            
        except Exception as e:
            error_msg = str(e)
            detailed_msg = (
                f'生成文档时出错：{error_msg}\n\n'
                f'请检查：\n'
                f'1. Excel单元格引用是否正确\n'
                f'2. 映射关系是否配置正确\n'
                f'3. Word模板中的变量名是否正确\n'
                f'4. Excel中是否存在空值'
            )
            QMessageBox.critical(self, '错误', detailed_msg)

    def _process_document(self, doc, sheet, row_num, mappings):
        """处理单个文档的变量替换"""
        try:
            print(f"处理数据行: {row_num}")  # 添加调试信息
            # 处理正文中的变量
            for paragraph in doc.paragraphs:
                text = paragraph.text
                original_text = text
                for word_var, cell_ref in mappings.items():
                    if word_var in text:
                        try:
                            # 从单元格引用中提取列字母
                            col_letter = ''.join(c for c in cell_ref if c.isalpha())
                            # 使用当前数据行号和列字母构建新的单元格引用
                            current_cell_ref = f"{col_letter}{row_num}"
                            value = sheet[current_cell_ref].value
                            if value is None:
                                value = ''  # 如果单元格为空，替换为空字符串
                            text = text.replace(word_var, str(value))
                        except Exception as e:
                            text = text.replace(word_var, '')  # 如果获取值出错，也替换为空字符串
                            print(f"替换变量出错: {word_var} -> {e}")  # 添加调试信息
                if text != original_text:  # 只有当文本发生变化时才更新
                    paragraph.text = text
                    print(f"更新段落: {original_text} -> {text}")  # 添加调试信息
            
            # 处理表格中的变量
            for table in doc.tables:
                for row_cells in table.rows:
                    for cell in row_cells.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            original_text = text
                            for word_var, cell_ref in mappings.items():
                                if word_var in text:
                                    try:
                                        # 从单元格引用中提取列字母
                                        col_letter = ''.join(c for c in cell_ref if c.isalpha())
                                        # 使用当前数据行号和列字母构建新的单元格引用
                                        current_cell_ref = f"{col_letter}{row_num}"
                                        value = sheet[current_cell_ref].value
                                        if value is None:
                                            value = ''
                                        text = text.replace(word_var, str(value))
                                    
                                    except Exception as e:
                                        text = text.replace(word_var, '')
                                        print(f"替换变量出错: {word_var} -> {e}")  # 添加调试信息
                            if text != original_text:
                                paragraph.text = text
        except Exception as e:
            raise Exception(f'处理文档时出错：{str(e)}')

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = ExcelWordConverter()
    window.show()
    sys.exit(app.exec()) 